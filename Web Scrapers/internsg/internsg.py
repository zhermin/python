from urllib.request import Request, urlopen
from bs4 import BeautifulSoup as soup
from docx import Document
from time import sleep
from random import randint


while True:

	doc = int(input("[1] Excel\n[2] Word\n"))

	if doc == 1:

		filename = "internsg_excel.csv"
		f = open(filename, "w")

		headers = "No., Company, Position, Requirements, Description\n"
		f.write(headers)
		break

	elif doc == 2:

		filename = "internsg-word.docx"
		document = Document()
		document.add_heading("InternSG", 0)
		break

	else:
		continue


def parse_html(current_url):

	req = Request(current_url, headers={'User-Agent': 'Mozilla/5.0'})
	webpage = urlopen(req).read()
	urlopen(req).close()

	page_soup = soup(webpage, "html.parser")

	[s.decompose() for s in page_soup('script')]

	return page_soup

def configure_items(current_container, all_headers, page_no, num, doc):

	for container in range (len(current_container)):			

		print("NUMBER {}".format(num))

		if doc == 2:
			sleep(randint(2,5))

		header_items = []
		header_items.append(num) # append [0]
		num += 1

		all_info = current_container[container].text[1:-3].replace(",", " -").replace("\n", ", ")
		the_link = current_container[container].a["href"]

		info_list = all_info.split(", ")

		for i in range (len(info_list)): # append [1 to 5]
			header_items.append(info_list[i])
		

		if doc == 1: # if excel

			for excel_items in range (3):
				header_items.append([info_list[excel_items]])
				break

		# if word

		header_items.append(the_link) # append [6]
		link_soup = parse_html(the_link) # access link of the current item


		company_profile_list = []
		company_profile = link_soup.findAll("dl", {"class":"dl-horizontal even"})
		company_profile_text = company_profile[5].dd.findAll()

		for ii in range (len(company_profile_text)):
			company_profile_list.append(company_profile_text[ii].text)

		# print(company_profile_list)
		header_items.append(company_profile_list) # append [7]


		job_description_list = []
		job_description = link_soup.findAll("dl", {"class":"dl-horizontal odd"})
		job_description_text = job_description[-1].dd.findAll()

		for iii in range (len(job_description_text)):
			job_description_list.append(job_description_text[iii].text)

		# print(job_description_list)
		header_items.append(job_description_list) # append [8]

		all_headers.append(header_items)
	
	return all_headers, num



page_no = 0
num = 1

while True:

	page_no += 1
	current_url = "https://www.internsg.com/jobs/go/{}/?type=job&filter_tax_job_profession=107&filter_tax_job_industry=0&filter_s=".format(str(page_no))

	sleep(randint(2,3))
	print("\nPAGE {}".format(page_no))

	page_soup = parse_html(current_url)

	odd_containers = page_soup.findAll("div", {"class":"row row-job odd"})
	even_containers = page_soup.findAll("div", {"class":"row row-job even"})
	all_containers = [odd_containers, even_containers]

	if len(odd_containers) == 0 or len(even_containers) == 0:
		break

	all_headers = []
	header_list = ["No.", "Company", "Position", "Timeline", "Location", "Posted", "Link", "Company Profile", "Job Description"]

	for all_container in range (len(all_containers)):
		current_container = all_containers[all_container]
		all_headers, num = configure_items(current_container, all_headers, page_no, num, doc)

	for x in range (len(all_headers)):

		for xx in range (len(all_headers[x])):

			if doc == 1:

				if xx == 3:
					break

				f.write(str(all_headers[x][xx]) + ", ")
				continue


			current_item = all_headers[x][xx]

			if xx == 7 or xx == 8:
				document.add_paragraph(header_list[xx] + " : {}".format("\n".join(current_item)))
				continue

			document.add_paragraph(header_list[xx] + " : " + str(current_item))

		if doc == 1:
			f.write("\n")
			continue

		document.add_page_break()




print("\nD O N E")

if doc == 1:
	f.close()
else:
	document.save(filename)



# IMPORTANT !!!

# REMEMBER TO CLOSE THE FILE YOU ARE WORKING ON
# OTHERWISE, THERE WILL BE AN ERROR WHEN OPENING THE CSV FILE
# ERROR : THIS FILE IS LOCKED FOR EDITING BY 'ANOTHER USER'

# LESSON LEARNT
# DON'T HAVE USER INPUT() BEFORE CLOSING THE FILE
# OTHERWISE, THE FILE WON'T CLOSE PROPERLY